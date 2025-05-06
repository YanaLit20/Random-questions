from docx import Document
from docx.shared import Inches
import random

questions = Document('вопросы.docx')
paragraphs = questions.paragraphs
n = list(range(0, len(paragraphs)+1))
random.shuffle(n)

doc = Document()
doc._body.clear_content() 
k = 1

while k < 21:
    a = paragraphs[k]
    doc.add_paragraph(a.text)
    k += 1


doc.save('question.docx')